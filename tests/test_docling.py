"""Docling integration: real headings, tables as markdown, new formats.

The docling extra is heavy and NOT installed in CI — the whole module
skips when docling is missing. The legacy extractors keep their own
coverage; these tests pin the docling path (tables survive chunking,
headings feed the TOC, xlsx/images become loadable).
"""

import os
import subprocess
import sys
from pathlib import Path

import pytest

pytest.importorskip("docling")

from rag_kit import _docling  # noqa: E402
from rag_kit._docling import extract_document, is_available  # noqa: E402
from rag_kit._rag import RAGSystem  # noqa: E402


@pytest.fixture(scope="module")
def docx_path(tmp_path_factory):
    from docx import Document

    path = tmp_path_factory.mktemp("docs") / "sample.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("The quick brown fox jumps over the lazy dog. " * 4)
    doc.add_heading("Setup Guide", level=2)
    doc.add_paragraph("First install the package, then configure the API key.")
    tbl = doc.add_table(rows=2, cols=3)
    tbl.rows[0].cells[0].text = "Part"
    tbl.rows[0].cells[1].text = "Rating"
    tbl.rows[0].cells[2].text = "Price"
    tbl.rows[1].cells[0].text = "A-100"
    tbl.rows[1].cells[1].text = "12V 3A"
    tbl.rows[1].cells[2].text = "$4.50"
    doc.save(str(path))
    return str(path)


@pytest.fixture(scope="module")
def xlsx_path(tmp_path_factory):
    from openpyxl import Workbook

    path = tmp_path_factory.mktemp("docs") / "prices.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Prices"
    ws.append(["Item", "Cost"])
    ws.append(["Bolt", "0.10"])
    ws.append(["Nut", "0.05"])
    wb.save(str(path))
    return str(path)


@pytest.fixture(scope="module")
def html_path(tmp_path_factory):
    path = tmp_path_factory.mktemp("docs") / "page.html"
    path.write_text(
        "<html><body><h1>Widget Manual</h1><p>Welcome to the widget manual.</p>"
        "<h2>Assembly</h2><p>Attach part A to part B.</p>"
        "<table><tr><th>Part</th><th>Torque</th></tr>"
        "<tr><td>Bolt</td><td>5 Nm</td></tr></table></body></html>"
    )
    return str(path)


def test_is_available():
    assert is_available()


def test_cuda_ep_usable_returns_bool():
    # Environment-dependent result, but must never raise and always be a
    # definite bool (False on CPU-only ort, on load failure, or no GPU).
    assert isinstance(_docling.cuda_ep_usable(), bool)


def test_extract_docx_keeps_tables_and_headings(docx_path):
    out = extract_document(docx_path)
    # Table content survives as a markdown table.
    assert "12V 3A" in out["text"]
    assert "| Part" in out["text"]
    # Real headings with hierarchy and offsets into the returned text.
    titles = [h["title"] for h in out["headings"]]
    assert "Introduction" in titles
    assert "Setup Guide" in titles
    levels = {h["title"]: h["level"] for h in out["headings"]}
    assert levels["Introduction"] < levels["Setup Guide"]  # h1 < h2
    offsets = [h["offset"] for h in out["headings"]]
    assert offsets == sorted(offsets)  # reading order
    for h in out["headings"]:
        assert out["text"][h["offset"] :].startswith("#" * h["level"] + " ")


def test_extract_xlsx_renders_table(xlsx_path):
    out = extract_document(xlsx_path)
    assert "Bolt" in out["text"]
    assert "0.10" in out["text"]


def test_extract_html_headings(html_path):
    out = extract_document(html_path)
    titles = [h["title"] for h in out["headings"]]
    assert titles == ["Widget Manual", "Assembly"]  # h1 then h2


@pytest.fixture(scope="module")
def vtt_path(tmp_path_factory):
    path = tmp_path_factory.mktemp("docs") / "captions.vtt"
    path.write_text(
        "WEBVTT\n\n"
        "00:00:01.000 --> 00:00:04.000\n"
        "The quick brown fox jumps over the lazy dog.\n\n"
        "00:00:05.000 --> 00:00:08.000\n"
        "Setup Guide: install the package first.\n"
    )
    return str(path)


def test_extract_vtt_subtitles(vtt_path):
    # VTT needs no ASR extra — docling parses the cues directly.
    out = extract_document(vtt_path)
    assert "quick brown fox" in out["text"].lower()


def test_audio_pipeline_runs(tmp_path):
    # Smoke test only: whisper on a 440 Hz tone produces no words, but
    # the docling audio pipeline (ASR) must run without raising.
    pytest.importorskip("whisper")
    import wave

    import numpy as np

    path = tmp_path / "tone.wav"
    samples = (0.1 * np.sin(2 * np.pi * 440 * np.arange(16000) / 16000)).astype(np.float32)
    with wave.open(str(path), "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(16000)
        w.writeframes((samples * 32767).astype(np.int16).tobytes())
    out = extract_document(str(path))
    assert "text" in out
    assert _docling.source_type_for(".wav") == "audio"


def _runner_cmd() -> list[str]:
    root = Path(__file__).resolve().parents[1]
    return [sys.executable, str(root / "scripts" / "docling_convert.py")]


def _runner_env() -> dict:
    root = Path(__file__).resolve().parents[1]
    return {**os.environ, "PYTHONPATH": str(root / "src")}


def test_docling_convert_runner_writes_markdown(docx_path, tmp_path):
    out_dir = tmp_path / "out"
    r = subprocess.run(
        [*_runner_cmd(), docx_path, "--out-dir", str(out_dir)],
        capture_output=True,
        text=True,
        env=_runner_env(),
        timeout=300,
    )
    assert r.returncode == 0, r.stderr
    md = (out_dir / "sample.md").read_text(encoding="utf-8")
    assert "12V 3A" in md  # docling table survived the round-trip


def test_docling_convert_runner_check_gpu():
    r = subprocess.run(
        [*_runner_cmd(), "--check-gpu"],
        capture_output=True,
        text=True,
        env=_runner_env(),
        timeout=120,
    )
    assert r.returncode == 0, r.stderr
    assert "CUDA EP usable" in r.stdout


def test_load_file_docx_uses_docling_path(docx_path, tmp_path):
    rag = RAGSystem(db_path=str(tmp_path / "dl.db"))
    fid = rag.load_file(docx_path)
    info = rag.list()[0]
    assert info["source_type"] == "docx"
    # Table rows survived chunking — only docling produces this; the
    # legacy python-docx path reads paragraphs only.
    chunk_text = " ".join((rag.get_chunk(fid, i) or {}).get("text", "") for i in range(20))
    assert "12V 3A" in chunk_text
    # Docling headings feed the TOC machinery (regex heuristic would
    # never detect "Setup Guide" — no number prefix, not ALL CAPS).
    titles = [m["title"] for m in rag._storage.get_section_mappings(fid)]
    assert "Setup Guide" in titles


def test_load_file_xlsx_new_format(xlsx_path, tmp_path):
    rag = RAGSystem(db_path=str(tmp_path / "dl.xlsx.db"))
    fid = rag.load_file(xlsx_path)
    chunk_text = " ".join((rag.get_chunk(fid, i) or {}).get("text", "") for i in range(10))
    assert "Bolt" in chunk_text


def test_pdf_empty_from_docling_falls_back_to_legacy(tmp_path, monkeypatch):
    # A blank PDF yields empty docling text → the legacy pypdf chain must
    # still run (scanned/blank docs keep working).
    from pypdf import PdfWriter

    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(str(path))

    import rag_kit._rag as rag_mod

    monkeypatch.setattr(rag_mod, "_has_meaningful_text", lambda text: True)
    rag = RAGSystem(db_path=str(tmp_path / "fallback.db"))
    rag.load_file(str(path))
    assert rag.list()[0]["source_type"] == "pdf"


def test_fallback_when_docling_missing(docx_path, xlsx_path, tmp_path, monkeypatch):
    monkeypatch.setattr(_docling, "is_available", lambda: False)
    rag = RAGSystem(db_path=str(tmp_path / "legacy.db"))
    # DOCX still loads via python-docx.
    rag.load_file(docx_path)
    assert rag.list()[0]["source_type"] == "docx"
    # Docling-only formats must fail with an actionable install hint.
    with pytest.raises(ImportError, match="docling"):
        rag.load_file(xlsx_path)


def test_docling_opt_out_env(monkeypatch):
    # RAGKIT_DOCLING=0 disables the docling path even when installed.
    monkeypatch.setenv("RAGKIT_DOCLING", "0")
    assert not is_available()
    monkeypatch.delenv("RAGKIT_DOCLING")
    assert is_available()
